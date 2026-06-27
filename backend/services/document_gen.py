"""
services/document_gen.py - Generate downloadable documents

Creates:
- DOCX files for cover letters (using python-docx)
- PDF files for resumes (using reportlab or weasyprint)
"""

import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_cover_letter_docx(
    cover_letter_content: str,
    candidate_name: str,
    company: str
) -> bytes:
    """
    Generate a formatted DOCX cover letter.
    
    Args:
        cover_letter_content: Markdown content of cover letter
        candidate_name: Name for the header
        company: Company name (for filename suggestion)
    
    Returns:
        DOCX file as bytes
    """
    doc = Document()
    
    # Set up the document
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Parse markdown content and add to document
    # This is a simple parser - handles basic markdown
    lines = cover_letter_content.split("\n")
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.add_run(line[2:]).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.add_run(line[3:]).bold = True
        # Handle regular text
        else:
            # Handle bold (**text**)
            p = doc.add_paragraph()
            
            # Simple bold handling
            import re
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def generate_resume_docx(
    resume_content: str,
    candidate_name: str
) -> bytes:
    """
    Generate a formatted DOCX resume.
    
    Similar to cover letter but with different formatting
    appropriate for resumes.
    """
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    lines = resume_content.split("\n")
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Handle headers
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(12)
            p.space_before = Pt(12)
        # Handle bullet points
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
